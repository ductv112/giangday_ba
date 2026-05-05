"""Shared docx helpers cho tất cả TPL builders."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

BRAND_RED = RGBColor(0xD4, 0x1B, 0x3F)
BRAND_DARK = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_GREY = RGBColor(0x4A, 0x4A, 0x4A)
PLACEHOLDER_GREY = RGBColor(0x99, 0x99, 0x99)
EMERALD = RGBColor(0x05, 0x96, 0x69)

OUT_DIR = Path(__file__).parent.parent / 'bieu-mau-docx'


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def new_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    return doc


def add_brand_header(doc, tpl_id, buoi_label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('iPMAC ')
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = BRAND_RED
    r = p.add_run(f'· Khoá Ready for BA · {tpl_id}{(" · " + buoi_label) if buoi_label else ""}')
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.color.rgb = TEXT_GREY


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = BRAND_RED
    if subtitle:
        sp = doc.add_paragraph()
        sr = sp.add_run(subtitle)
        sr.font.name = 'Calibri'
        sr.font.size = Pt(11)
        sr.italic = True
        sr.font.color.rgb = TEXT_GREY


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = BRAND_RED
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = BRAND_DARK
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = BRAND_DARK
    return p


def add_para(doc, text, italic=False, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    if italic:
        r.italic = True
    if bold:
        r.font.bold = True
    if color:
        r.font.color.rgb = color
    return p


def add_placeholder(doc, hint):
    p = doc.add_paragraph()
    r = p.add_run(f'[ {hint} ]')
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = PLACEHOLDER_GREY
    return p


def add_bullet(doc, text, italic=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        if italic:
            r.italic = True
            r.font.color.rgb = PLACEHOLDER_GREY
    if not p.runs:
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        if italic:
            r.italic = True
            r.font.color.rgb = PLACEHOLDER_GREY
    else:
        # set text via runs
        p.runs[0].text = text
    return p


def add_checkbox(doc, text):
    p = doc.add_paragraph()
    r = p.add_run('☐ ')
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    return p


def add_table_with_headers(doc, headers, rows, header_bg='D41B3F', header_white=True, italic_placeholder=True):
    """Create a table with red header + sample/empty rows.

    rows: list of tuples — empty string for blank fill cell.
    """
    n_rows = len(rows) + 1
    n_cols = len(headers)
    t = doc.add_table(rows=n_rows, cols=n_cols)
    t.style = 'Light Grid Accent 1'
    # Header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        if header_white:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, header_bg)
    # Data
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            if italic_placeholder and (str(val).startswith('[') or str(val).startswith('—')):
                r.italic = True
                r.font.color.rgb = PLACEHOLDER_GREY
    return t


def add_signoff_table(doc, roles=None):
    if roles is None:
        roles = ['BA', 'Tech Lead', 'Sponsor', 'QA Lead']
    headers = ['Vai trò', 'Tên', 'Chữ ký', 'Ngày']
    rows = [(role, '', '', '') for role in roles]
    return add_table_with_headers(doc, headers, rows)


def add_meta_kv_table(doc, kv_pairs):
    """2-col Key-Value table với label đậm + value italic placeholder."""
    t = doc.add_table(rows=len(kv_pairs), cols=2)
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    t.columns[0].width = Cm(4.5)
    t.columns[1].width = Cm(12.5)
    for i, (k, v) in enumerate(kv_pairs):
        c1 = t.rows[i].cells[0]
        c2 = t.rows[i].cells[1]
        c1.text = ''
        c2.text = ''
        r1 = c1.paragraphs[0].add_run(k)
        r1.font.bold = True
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        r2 = c2.paragraphs[0].add_run(v)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(11)
        if v.startswith('[') or v.startswith('☐'):
            r2.italic = True
            r2.font.color.rgb = PLACEHOLDER_GREY
        set_cell_bg(c1, 'F5F5F5')
    return t


def add_footer_note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = TEXT_GREY


def save(doc, filename):
    OUT_DIR.mkdir(exist_ok=True)
    out = OUT_DIR / filename
    doc.save(str(out))
    return out
