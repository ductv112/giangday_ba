"""Build TPL-017 SRS template as proper Word form (python-docx).

Demo: 1 file → user xem chất lượng → nếu OK thì batch toàn bộ.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# iPMAC brand colors
BRAND_RED = RGBColor(0xD4, 0x1B, 0x3F)
BRAND_DARK = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_GREY = RGBColor(0x4A, 0x4A, 0x4A)
PLACEHOLDER_GREY = RGBColor(0x99, 0x99, 0x99)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = BRAND_RED
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK
    return p


def add_para(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    return p


def add_placeholder(doc, hint):
    """Italic grey placeholder text — học viên fill thay."""
    p = doc.add_paragraph()
    run = p.add_run(f'[ {hint} ]')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = PLACEHOLDER_GREY
    return p


def add_metadata_table(doc):
    rows = [
        ('Project', '[ Điền tên dự án ]'),
        ('Module', '[ Tên phân hệ ]'),
        ('Version', '1.0'),
        ('Status', '☐ Draft   ☐ Reviewed   ☐ Approved'),
        ('BA', '[ Tên BA ]'),
        ('Date', '__/__/____'),
        ('Approved by', '[ Sponsor name + ngày ký ]'),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    t.columns[0].width = Cm(4)
    t.columns[1].width = Cm(13)
    for i, (k, v) in enumerate(rows):
        c1 = t.rows[i].cells[0]
        c2 = t.rows[i].cells[1]
        c1.text = ''
        c2.text = ''
        c1.paragraphs[0].add_run(k).font.bold = True
        c2.paragraphs[0].add_run(v).italic = '[' in v
        for c in (c1, c2):
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(11)
        set_cell_bg(c1, 'F5F5F5')
    return t


def add_change_log_table(doc):
    headers = ['Version', 'Ngày', 'Người', 'Thay đổi']
    t = doc.add_table(rows=3, cols=4)
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        run = c.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.name = 'Calibri'
        set_cell_bg(c, 'D41B3F')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # 2 sample rows
    samples = [
        ('1.0', '__/__/____', '[ BA ]', 'Bản đầu'),
        ('', '', '', ''),
    ]
    for i, row in enumerate(samples, 1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ''
            r = c.paragraphs[0].add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            if val.startswith('['):
                r.italic = True
                r.font.color.rgb = PLACEHOLDER_GREY
    return t


def add_user_story_table(doc):
    headers = ['ID', 'User Story', 'Acceptance Criteria', 'Priority']
    t = doc.add_table(rows=3, cols=4)
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.name = 'Calibri'
        set_cell_bg(c, 'D41B3F')
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sample_rows = [
        ('US-001', 'Là [vai trò], tôi muốn [tính năng], để [lợi ích]',
         'Given... When... Then...', 'MUST'),
        ('US-002', '', '', ''),
    ]
    for i, row in enumerate(sample_rows, 1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ''
            r = c.paragraphs[0].add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            if val and j > 0 and i == 1:
                r.italic = True
                r.font.color.rgb = PLACEHOLDER_GREY
    return t


def add_fr_table(doc):
    headers = ['ID', 'Mô tả', 'Source', 'Priority', 'Status']
    t = doc.add_table(rows=3, cols=5)
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.name = 'Calibri'
        set_cell_bg(c, 'D41B3F')
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    samples = [
        ('FR-001', 'Hệ thống PHẢI ...', 'UC-001', 'MUST', 'Approved'),
        ('FR-002', '', '', '', ''),
    ]
    for i, row in enumerate(samples, 1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ''
            r = c.paragraphs[0].add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            if val and i == 1 and j > 0 and 'PHẢI' in val:
                r.italic = True
                r.font.color.rgb = PLACEHOLDER_GREY
    return t


def add_nfr_table(doc, nfr_type, items):
    headers = ['ID', 'Yêu cầu', 'Mục tiêu']
    t = doc.add_table(rows=len(items) + 1, cols=3)
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.name = 'Calibri'
        set_cell_bg(c, 'F5F5F5')
    for i, (id_, req, target) in enumerate(items, 1):
        for j, val in enumerate([id_, req, target]):
            c = t.rows[i].cells[j]
            c.text = ''
            r = c.paragraphs[0].add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
    return t


def add_signoff_table(doc):
    headers = ['Vai trò', 'Tên', 'Chữ ký', 'Ngày']
    rows = [('BA', '', '', ''), ('Tech Lead', '', '', ''),
            ('Sponsor', '', '', ''), ('QA Lead', '', '', '')]
    t = doc.add_table(rows=len(rows) + 1, cols=4)
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.name = 'Calibri'
        set_cell_bg(c, 'D41B3F')
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ''
            r = c.paragraphs[0].add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
    return t


# ============================================================
def build():
    doc = Document()

    # Page setup — A4 portrait
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # === HEADER LOGO + Title ===
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_p.add_run('iPMAC ')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = BRAND_RED
    run = header_p.add_run('· Khoá Ready for BA · TPL-017')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_GREY

    # H1 Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(12)
    run = title_p.add_run('SRS — Software Requirements Specification')
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = BRAND_RED

    add_para(
        doc,
        'Template chuẩn 8 mục — copy về điền cho mọi dự án. '
        'Pass-rate ≥ 70/100 theo Rubric (TPL-033).',
        italic=True, color=TEXT_GREY
    )

    # === METADATA ===
    add_h2(doc, 'Metadata')
    add_metadata_table(doc)

    # === CHANGE LOG ===
    add_h2(doc, 'Change Log')
    add_change_log_table(doc)

    doc.add_page_break()

    # === 1. GIỚI THIỆU ===
    add_h1(doc, '1. Giới thiệu')

    add_h2(doc, '1.1. Mục đích tài liệu')
    add_placeholder(doc, 'Tài liệu mô tả chi tiết yêu cầu cho phân hệ X — input cho dev/QA/designer')

    add_h2(doc, '1.2. Phạm vi')
    add_h3(doc, 'In-scope (PHẢI làm)')
    add_placeholder(doc, 'Liệt kê chức năng / phân hệ / module cần làm')
    add_h3(doc, 'Out-of-scope (KHÔNG làm)')
    add_placeholder(doc, 'Liệt kê những gì KHÔNG nằm trong phase này')

    add_h2(doc, '1.3. Định nghĩa & viết tắt (Glossary)')
    add_para(doc, '(Tham chiếu TPL-020 hoặc liệt kê trực tiếp dưới đây)')
    glossary = doc.add_table(rows=4, cols=2)
    glossary.style = 'Light Grid Accent 1'
    glossary.rows[0].cells[0].text = 'Thuật ngữ'
    glossary.rows[0].cells[1].text = 'Định nghĩa'
    for c in glossary.rows[0].cells:
        for r in c.paragraphs[0].runs:
            r.font.bold = True
        set_cell_bg(c, 'F5F5F5')
    glossary.rows[1].cells[0].text = 'OTP'
    glossary.rows[1].cells[1].text = 'One-Time Password — mã 6 số có thời hạn 5 phút'

    add_h2(doc, '1.4. Tài liệu tham chiếu')
    for ref in ['BRD: [link]', 'Mockup Figma: [link]', 'API spec: [link]']:
        p = doc.add_paragraph(ref, style='List Bullet')
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

    doc.add_page_break()

    # === 2. MÔ TẢ TỔNG QUAN ===
    add_h1(doc, '2. Mô tả tổng quan')

    add_h2(doc, '2.1. Bối cảnh')
    add_placeholder(doc, '2-3 đoạn: Vì sao có dự án? Hiện trạng AS-IS? Mục tiêu TO-BE?')

    add_h2(doc, '2.2. Chức năng chính (high-level)')
    add_placeholder(doc, 'Liệt kê 3-7 chức năng chính bằng bullet')

    add_h2(doc, '2.3. Người dùng')
    user_t = doc.add_table(rows=3, cols=3)
    user_t.style = 'Light Grid Accent 1'
    headers = ['Role', 'Mô tả', 'Số lượng dự kiến']
    for i, h in enumerate(headers):
        c = user_t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        set_cell_bg(c, 'F5F5F5')
    samples = [
        ('Khách hàng', 'End user dùng app', '~5,000'),
        ('Admin', 'Quản trị backoffice', '5'),
    ]
    for i, row in enumerate(samples, 1):
        for j, val in enumerate(row):
            user_t.rows[i].cells[j].text = ''
            r = user_t.rows[i].cells[j].paragraphs[0].add_run(val)
            r.font.italic = True
            r.font.color.rgb = PLACEHOLDER_GREY

    add_h2(doc, '2.4. Ràng buộc & giả định')
    add_placeholder(doc, 'Ràng buộc: phải tích hợp với hệ thống X')
    add_placeholder(doc, 'Giả định: người dùng có smartphone + 4G')

    doc.add_page_break()

    # === 3. YÊU CẦU CHỨC NĂNG ===
    add_h1(doc, '3. Yêu cầu chức năng (Functional Requirements)')

    add_h2(doc, '3.1. Use Case Diagram')
    add_placeholder(doc, '[ Đính kèm hình UC Diagram — link Figma hoặc paste ảnh ]')

    add_h2(doc, '3.2. Use Case Details')
    add_placeholder(doc, 'Liệt kê link UC Spec — sử dụng template TPL-013')

    add_h2(doc, '3.3. User Stories')
    add_user_story_table(doc)

    add_h2(doc, '3.4. Functional Requirements (chi tiết)')
    add_fr_table(doc)

    doc.add_page_break()

    # === 4. NFR ===
    add_h1(doc, '4. Yêu cầu phi chức năng (Non-functional Requirements)')

    add_h2(doc, '4.1. Performance')
    add_nfr_table(doc, 'Performance', [
        ('NFR-001', 'Thời gian load trang chủ', '≤ 2 giây trên 4G'),
        ('NFR-002', '', ''),
        ('NFR-003', '', ''),
    ])

    add_h2(doc, '4.2. Security')
    add_nfr_table(doc, 'Security', [
        ('NFR-010', 'Mọi password PHẢI hash bcrypt cost 12', ''),
        ('NFR-011', '', ''),
    ])

    add_h2(doc, '4.3. Usability')
    add_nfr_table(doc, 'Usability', [
        ('NFR-020', 'Mobile app PHẢI tuân thủ Material Design / iOS HIG', ''),
        ('NFR-021', '', ''),
    ])

    add_h2(doc, '4.4. Reliability')
    add_nfr_table(doc, 'Reliability', [
        ('NFR-030', 'Uptime ≥ 99.5%', '≤ 3.6 giờ downtime/tháng'),
        ('NFR-031', '', ''),
    ])

    doc.add_page_break()

    # === 5-7 (placeholder ngắn) ===
    add_h1(doc, '5. Mô hình dữ liệu')
    add_h2(doc, '5.1. ERD')
    add_placeholder(doc, '[ Đính kèm hình ERD — sử dụng TPL-012 ]')
    add_h2(doc, '5.2. Định nghĩa Entity')
    add_placeholder(doc, 'Liệt kê chi tiết Entity + Attribute (table cho mỗi entity)')

    add_h1(doc, '6. Giao diện người dùng')
    add_h2(doc, '6.1. Wireframe / Mockup')
    add_placeholder(doc, 'Bảng liệt kê: Màn hình | Wireframe (link Figma) | Mockup HD')
    add_h2(doc, '6.2. Flow màn hình')
    add_placeholder(doc, '[ Đính kèm flowchart navigation ]')
    add_h2(doc, '6.3. Design System')
    for line in ['Color: primary, secondary', 'Font: Inter / Roboto',
                 'Component library: shadcn / AntD', 'Icon: Lucide']:
        p = doc.add_paragraph(line, style='List Bullet')
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

    add_h1(doc, '7. Tích hợp')
    add_h2(doc, '7.1. Hệ thống bên ngoài')
    integration_t = doc.add_table(rows=3, cols=4)
    integration_t.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Hệ thống', 'Mục đích', 'API endpoint', 'Auth']):
        c = integration_t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        set_cell_bg(c, 'F5F5F5')

    add_h2(doc, '7.2. API Contract')
    add_placeholder(doc, 'API spec link Postman / Swagger')

    doc.add_page_break()

    # === 8. PHỤ LỤC ===
    add_h1(doc, '8. Phụ lục')

    add_h2(doc, '8.1. Glossary (chi tiết)')
    add_para(doc, 'Sử dụng TPL-020 — Glossary template')

    add_h2(doc, '8.2. Traceability Matrix')
    add_para(doc, 'Sử dụng TPL-027 — Traceability template')

    add_h2(doc, '8.3. Test Cases')
    add_para(doc, 'Sử dụng TPL-019 — Test Cases template')

    add_h2(doc, '8.4. Approval (Sign-off)')
    add_signoff_table(doc)

    # Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(40)
    run = footer_p.add_run('— END OF SRS · Version 1.0 —')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = TEXT_GREY

    out = Path(__file__).parent.parent / 'bieu-mau-docx' / 'TPL-017 — SRS Đặc tả phần mềm (Buổi 7).docx'
    out.parent.mkdir(exist_ok=True)
    doc.save(str(out))
    print(f'OK Saved: {out.name} ({out.stat().st_size} bytes)')


if __name__ == '__main__':
    build()
